from docx import Document
from io import BytesIO
from typing import List

class DOCXExtractor:
    @staticmethod
    def extract_text(file_content: bytes) -> str:
        """Extract text from DOCX file"""
        try:
            doc_file = BytesIO(file_content)
            doc = Document(doc_file)
            
            text = []
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    text.append(paragraph.text)
            
            # Extract text from tables
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        if cell.text.strip():
                            text.append(cell.text)
            
            return "\n".join(text)
        except Exception as e:
            raise Exception(f"Error extracting DOCX: {str(e)}")
    
    @staticmethod
    def check_formatting_issues(file_content: bytes) -> List[str]:
        """Check for ATS-unfriendly formatting in DOCX"""
        issues = []
        try:
            doc_file = BytesIO(file_content)
            doc = Document(doc_file)
            
            # Check for text boxes
            for shape in doc.inline_shapes:
                issues.append("Contains text boxes or embedded objects - avoid for ATS")
                break
            
            # Check for headers/footers
            if any(section.header.paragraphs for section in doc.sections):
                if any(p.text.strip() for section in doc.sections for p in section.header.paragraphs):
                    issues.append("Important info in header - ATS may not read it")
            
            # Check for tables (excessive use)
            if len(doc.tables) > 2:
                issues.append("Multiple tables detected - simplify structure for ATS")
            
            return issues
        except:
            return []